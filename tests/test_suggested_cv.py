"""Tests for suggested-CV layout contract, word trim, and field merge."""

from pathlib import Path

from docx import Document

from backend.app.services.cv_layout_contract import (
    attach_docx_targets,
    parse_slot_layout_from_excerpt,
)
from backend.app.services.tailor_cv_llm import (
    apply_word_trim_fallback,
    lock_passing_fields,
    merge_with_locked,
    validate_generated_slots,
)
from backend.app.services.word_trim import word_trim_to_max


def test_parse_slot_layout_per_line_budgets():
    excerpt = (
        "JobPilot (LangGraph, FastAPI)\n"
        "- Built multi-agent search orchestration\n"
        "- Retrieved project evidence with RAG"
    )
    layout = parse_slot_layout_from_excerpt(0, excerpt)
    assert layout["title"]["max_characters"] == len("JobPilot (LangGraph, FastAPI)")
    assert len(layout["description_items"]) == 2
    assert layout["description_items"][0]["type"] == "bullet"
    assert layout["description_items"][0]["max_characters"] == len(
        "Built multi-agent search orchestration"
    )
    assert layout["description_items"][1]["max_characters"] == len(
        "Retrieved project evidence with RAG"
    )


def test_word_trim_drops_whole_words():
    text = "Built LangGraph multi-agent job search pipelines"
    trimmed = word_trim_to_max(text, 40)
    assert len(trimmed) <= 40
    assert "pipelines" not in trimmed or len(text) <= 40
    assert not trimmed.endswith("pipeli")


def test_next_suggested_cv_filename_increments(tmp_path, monkeypatch):
    from backend.app.services import suggested_cv_store as store

    monkeypatch.setattr(store, "drafts_dir", lambda _user_id: tmp_path)
    assert store.next_suggested_cv_filename(1, "Hamza_CV.docx") == "Hamza_CV_1.docx"
    (tmp_path / "Hamza_CV_1.docx").write_bytes(b"x")
    assert store.next_suggested_cv_filename(1, "Hamza_CV.docx") == "Hamza_CV_2.docx"
    (tmp_path / "Hamza_CV_2.docx").write_bytes(b"x")
    assert store.next_suggested_cv_filename(1, "Hamza_CV.docx") == "Hamza_CV_3.docx"


def test_lock_merge_preserves_good_fields():
    layout_slots = [
        {
            "slot_index": 0,
            "title": {"current_text": "A", "max_characters": 10},
            "description_items": [
                {
                    "item_index": 0,
                    "type": "bullet",
                    "current_text": "ok",
                    "max_characters": 20,
                },
                {
                    "item_index": 1,
                    "type": "bullet",
                    "current_text": "long",
                    "max_characters": 10,
                },
            ],
        }
    ]
    call1 = {
        "slots": [
            {
                "slot_index": 0,
                "title": "ShortTitle",
                "items": ["Fits in twenty chars", "This is way too long for ten"],
            }
        ]
    }
    locked = lock_passing_fields(call1, layout_slots)
    assert "0:title" in locked
    assert "0:items[0]" in locked
    assert "0:items[1]" not in locked

    call2 = {
        "slots": [
            {
                "slot_index": 0,
                "title": "WorsenedTitleXX",
                "items": ["Destroyed good line!!!", "Now short"],
            }
        ]
    }
    merged = merge_with_locked(call2, locked, layout_slots, [0])
    assert merged["slots"][0]["title"] == "ShortTitle"
    assert merged["slots"][0]["items"][0] == "Fits in twenty chars"
    assert merged["slots"][0]["items"][1] == "Now short"


def test_word_trim_fallback_and_validate():
    layout_slots = [
        {
            "slot_index": 1,
            "title": {"current_text": "Title", "max_characters": 8},
            "description_items": [
                {
                    "item_index": 0,
                    "type": "bullet",
                    "current_text": "abc",
                    "max_characters": 12,
                }
            ],
        }
    ]
    payload = {
        "slots": [
            {
                "slot_index": 1,
                "title": "TooLongTitle",
                "items": ["This bullet is oversized"],
            }
        ]
    }
    assert validate_generated_slots(payload, layout_slots, [1])
    trimmed, shortened, capped = apply_word_trim_fallback(payload, layout_slots)
    assert shortened is True
    assert capped
    assert validate_generated_slots(trimmed, layout_slots, [1]) == []


def test_attach_docx_targets(tmp_path: Path):
    path = tmp_path / "cv.docx"
    doc = Document()
    doc.add_paragraph("PROJECTS")
    doc.add_paragraph("JobPilot (LangGraph)")
    doc.add_paragraph("• Built the search agent")
    doc.add_paragraph("• Packaged fit analysis")
    doc.add_paragraph("SKILLS")
    doc.save(str(path))

    layout = parse_slot_layout_from_excerpt(
        0,
        "JobPilot (LangGraph)\n"
        "- Built the search agent\n"
        "- Packaged fit analysis",
    )
    enriched = attach_docx_targets([layout], path)[0]
    assert enriched["docx_targets"]["title_paragraph_index"] == 1
    assert enriched["docx_targets"]["item_paragraph_indexes"] == [2, 3]

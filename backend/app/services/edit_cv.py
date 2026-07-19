"""In-place .docx text replacement preserving paragraph/run styles."""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.run import Run


def _content_runs(runs: list[Run]) -> list[Run]:
    return [run for run in runs if (run.text or "").strip()]


def _style_run(runs: list[Run], *, prefer_bold: bool = False) -> Run:
    """Pick the run whose character style should apply to replacement text.

    CV titles often start with a non-bold whitespace run; the real title bold
    lives on a later run. Prefer a non-whitespace run, and when prefer_bold is
    set (project titles), prefer one that is already bold.
    """
    content = _content_runs(runs)
    candidates = content or list(runs)
    if prefer_bold:
        for run in candidates:
            if run.bold is True:
                return run
    return candidates[0]


def _copy_run_format(source: Run, target: Run) -> None:
    """Copy the character styles we care about for CV slot text."""
    if source is target:
        return
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    if source.font.size is not None:
        target.font.size = source.font.size
    if source.font.name:
        target.font.name = source.font.name


def _set_paragraph_text(
    paragraph,
    text: str,
    *,
    prefer_bold: bool = False,
) -> None:
    """Replace paragraph text while keeping the best content run's formatting."""
    runs = paragraph.runs
    if not runs:
        run = paragraph.add_run(text)
        if prefer_bold:
            run.bold = True
        return

    style_run = _style_run(runs, prefer_bold=prefer_bold)
    target = runs[0]
    _copy_run_format(style_run, target)
    if prefer_bold and style_run.bold is True:
        target.bold = True
    elif prefer_bold and any(run.bold is True for run in runs):
        # Title slots in this CV layout are bold even if the matched content
        # run had bold inherited as None in python-docx.
        target.bold = True
    target.text = text
    for run in runs[1:]:
        run.text = ""


def apply_slot_replacements(
    *,
    source_docx: Path,
    dest_docx: Path,
    layout_slots: list[dict[str, Any]],
    generated_slots: list[dict[str, Any]],
) -> None:
    """Copy master CV and write generated title/items into mapped paragraphs."""
    dest_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, dest_docx)
    doc = Document(str(dest_docx))
    paragraphs = list(doc.paragraphs)
    by_index = {int(slot["slot_index"]): slot for slot in generated_slots}

    for layout in layout_slots:
        slot_index = int(layout["slot_index"])
        generated = by_index.get(slot_index)
        if generated is None:
            continue
        targets = layout.get("docx_targets") or {}
        title_index = targets.get("title_paragraph_index")
        item_indexes = targets.get("item_paragraph_indexes") or []

        title = str(generated.get("title") or "")
        items = [str(item) for item in (generated.get("items") or [])]

        if title_index is not None and 0 <= int(title_index) < len(paragraphs):
            _set_paragraph_text(
                paragraphs[int(title_index)],
                title,
                prefer_bold=True,
            )

        for item_index, para_index in enumerate(item_indexes):
            if para_index is None:
                continue
            if item_index >= len(items):
                break
            idx = int(para_index)
            if 0 <= idx < len(paragraphs):
                _set_paragraph_text(paragraphs[idx], items[item_index])

    doc.save(str(dest_docx))
